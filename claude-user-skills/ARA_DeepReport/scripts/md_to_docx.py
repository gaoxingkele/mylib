"""通用 Markdown -> DOCX 转换器（中文排版 + 真超链接 + RAND 研究报告制作要素）。
用法: python md_to_docx.py <input.md> <output.docx> [--title "页眉短名"] [--version v1.3.0]

RAND 化要素（向后兼容：无下列标记的普通 MD 仍正常转换）：
- 页眉(报告短名) + 页脚(页码 + 版本)
- [[TOC]] 标记处插入 Word 自动目录字段（打开后右键"更新域"）
- ![alt](path.png) 图片嵌入（居中、按宽度缩放、alt 作图注）
- ::: 指令块：cover / keyfindings / recommendations / box [标题]
- 以"深挖/关键发现/政策启示/体系性/研判/时效提示"开头的引用块 → 自动底纹侧栏
"""
import re, sys, io, os, argparse
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CN_FONT = "宋体"
HEADING_FONT = "黑体"
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"(\*\*[^*]+\*\*)")
IMG_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
BOX_MARKERS = ("深挖", "关键发现", "政策启示", "体系性", "研判", "时效提示", "统一剧本")
BOX_FILLS = {"keyfindings": "DCE6F1", "recommendations": "E2EFDA", "box": "F2F2F2", "sidebar": "F2F2F2"}
BOX_TITLES = {"keyfindings": "关键发现", "recommendations": "政策启示与建议"}


def set_cn_font(run, font=CN_FONT, size_pt=11, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_hyperlink(paragraph, url, text, size_pt=11):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts"); rf.set(qn("w:eastAsia"), CN_FONT); rPr.append(rf)
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; r.append(t); h.append(r)
    paragraph._p.append(h)


def process_inline(p, text, font=CN_FONT, size_pt=11):
    idx = 0; segs = []
    for m in LINK_RE.finditer(text):
        if m.start() > idx:
            segs.append(("t", text[idx:m.start()]))
        segs.append(("l", (m.group(1), m.group(2)))); idx = m.end()
    if idx < len(text):
        segs.append(("t", text[idx:]))
    for kind, payload in segs:
        if kind == "l":
            add_hyperlink(p, payload[1], payload[0], size_pt)
        else:
            for part in BOLD_RE.split(payload):
                if not part:
                    continue
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2]); set_cn_font(r, font, size_pt, True)
                else:
                    r = p.add_run(part); set_cn_font(r, font, size_pt, False)


def add_heading_para(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    sizes = {1: 20, 2: 16, 3: 14, 4: 13, 5: 12, 6: 12}
    r = p.add_run(text.replace("**", "")); set_cn_font(r, HEADING_FONT, sizes.get(level, 12), True)
    return p


def add_table_from_md(doc, lines):
    rows = []
    for line in lines:
        if not line.strip():
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.match(r"^[-:\s]+$", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    n = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n); table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, ct in enumerate(row):
            if j < n:
                cell = table.cell(i, j); cell.text = ""
                process_inline(cell.paragraphs[0], ct, size_pt=9)
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_box(doc, title, body_lines, fill="F2F2F2"):
    table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"
    cell = table.cell(0, 0); _shade_cell(cell, fill); cell.text = ""
    first = True
    if title:
        p = cell.paragraphs[0]; r = p.add_run(title); set_cn_font(r, HEADING_FONT, 12, True)
        first = False
    for bl in body_lines:
        if not bl.strip():
            continue
        p = cell.paragraphs[0] if first else cell.add_paragraph(); first = False
        m = re.match(r"^\s*[-*]\s+(.+)$", bl)
        if m:
            p.paragraph_format.left_indent = Cm(0.5)
            process_inline(p, "• " + m.group(1))
        else:
            process_inline(p, bl)
    doc.add_paragraph()


def add_image(doc, path, alt):
    if not os.path.exists(path):
        p = doc.add_paragraph(); r = p.add_run(f"[缺图: {path}]"); set_cn_font(r, size_pt=9)
        return
    doc.add_picture(path, width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if alt:
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(alt); set_cn_font(r, size_pt=9); r.italic = True


def add_toc(doc):
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    t = OxmlElement("w:r"); tt = OxmlElement("w:t")
    tt.text = "（打开文档后右键此处选择“更新域”以生成目录）"
    t.append(tt); fld.append(t); p._p.append(fld)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin"); run._r.append(fc1)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "; run._r.append(it)
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end"); run._r.append(fc2)


def setup_header_footer(doc, title, version):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True  # 封面页不显示
    hdr = sec.header.paragraphs[0]; hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hdr.add_run(title); set_cn_font(r, font=CN_FONT, size_pt=9); r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    ftr = sec.footer.paragraphs[0]; ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(ftr)
    if version:
        r2 = ftr.add_run(f"    ·    {version}"); set_cn_font(r2, size_pt=9); r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_cover(doc, meta):
    for _ in range(3):
        doc.add_paragraph()
    if meta.get("series"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(meta["series"]); set_cn_font(r, HEADING_FONT, 12, True); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(24)
    r = p.add_run(meta.get("title", "")); set_cn_font(r, HEADING_FONT, 26, True)
    if meta.get("subtitle"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(8)
        r = p.add_run(meta["subtitle"]); set_cn_font(r, HEADING_FONT, 15, False); r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    for _ in range(6):
        doc.add_paragraph()
    for key in ("version", "date", "note"):
        if meta.get(key):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(meta[key]); set_cn_font(r, size_pt=11)
    pb = doc.add_paragraph(); run = pb.add_run(); br = OxmlElement("w:br"); br.set(qn("w:type"), "page"); run._r.append(br)


def md_to_docx(md_file, docx_file, title="", version=""):
    base_dir = os.path.dirname(os.path.abspath(md_file))
    lines = open(md_file, encoding="utf-8").read().split("\n")
    doc = Document()
    sec = doc.sections[0]
    for m in ("top", "bottom", "left", "right"):
        setattr(sec, f"{m}_margin", Cm(2.0))
    style = doc.styles["Normal"]; style.font.name = CN_FONT; style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if title:
        setup_header_footer(doc, title, version)

    i = 0; in_code = False; code_buf = []
    while i < len(lines):
        line = lines[i]
        mdir = re.match(r"^:::\s*(cover|keyfindings|recommendations|box|sidebar)\s*(.*)$", line.strip())
        if mdir and not in_code:
            name = mdir.group(1); arg = mdir.group(2).strip()
            buf = []; i += 1
            while i < len(lines) and not lines[i].strip().startswith(":::"):
                buf.append(lines[i]); i += 1
            i += 1
            if name == "cover":
                meta = {}
                for b in buf:
                    mm = re.match(r"^\s*(\w+)\s*[:：]\s*(.+)$", b)
                    if mm:
                        meta[mm.group(1).lower()] = mm.group(2).strip()
                add_cover(doc, meta)
            else:
                add_box(doc, arg or BOX_TITLES.get(name, ""), buf, BOX_FILLS.get(name, "F2F2F2"))
            continue
        if line.startswith("```"):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
                r = p.add_run("\n".join(code_buf)); r.font.name = "Consolas"; r.font.size = Pt(9)
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue
        if line.strip() == "[[TOC]]":
            add_toc(doc); i += 1; continue
        mimg = IMG_RE.match(line.strip())
        if mimg:
            path = mimg.group(2)
            if not os.path.isabs(path):
                path = os.path.join(base_dir, path)
            add_image(doc, path, mimg.group(1)); i += 1; continue
        if line.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl.append(lines[i]); i += 1
            add_table_from_md(doc, tbl); continue
        mh = re.match(r"^(#{1,6})\s+(.+)$", line)
        if mh:
            add_heading_para(doc, mh.group(2).strip(), len(mh.group(1))); i += 1; continue
        if line.strip() == "---":
            p = doc.add_paragraph(); r = p.add_run("—" * 30); set_cn_font(r, size_pt=10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; i += 1; continue
        if line.startswith(">"):
            qbuf = []
            while i < len(lines) and lines[i].startswith(">"):
                c = lines[i][1:]
                if c.startswith(" "):
                    c = c[1:]
                qbuf.append(c); i += 1
            head = qbuf[0] if qbuf else ""
            if any(mk in head for mk in BOX_MARKERS):
                add_box(doc, "", qbuf, "F2F2F2")
            else:
                for c in qbuf:
                    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.75)
                    process_inline(p, c)
            continue
        mb = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if mb:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.75 + len(mb.group(1)) * 0.4)
            process_inline(p, mb.group(2)); i += 1; continue
        mo = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        if mo:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75 + len(mo.group(1)) * 0.4)
            process_inline(p, f"{mo.group(2)}. {mo.group(3)}"); i += 1; continue
        if not line.strip():
            i += 1; continue
        p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.4; process_inline(p, line); i += 1

    doc.save(docx_file)
    print(f"Saved: {docx_file}")
    print(f"Size: {os.path.getsize(docx_file):,} bytes")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("md"); ap.add_argument("docx")
    ap.add_argument("--title", default=""); ap.add_argument("--version", default="")
    a = ap.parse_args()
    md_to_docx(a.md, a.docx, a.title, a.version)
